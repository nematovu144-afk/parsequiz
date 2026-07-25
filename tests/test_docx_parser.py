"""Unit tests for DocxParser (bold-run detection is what delimiter.py relies on)."""

import pytest
from docx import Document

from app.core.exceptions import CorruptedFileError
from app.parsers.docx_parser import DocxParser


def _make_docx(path):
    doc = Document()
    doc.add_paragraph("1. What color is the sky?")
    p = doc.add_paragraph("A) Red")
    p = doc.add_paragraph()
    run = p.add_run("B) Blue")
    run.bold = True
    doc.add_paragraph("C) Green")
    doc.save(str(path))


@pytest.mark.asyncio
async def test_extract_preserves_bold_run(tmp_path):
    path = tmp_path / "quiz.docx"
    _make_docx(path)

    paragraphs = await DocxParser().extract(path)

    texts = [p.plain_text for p in paragraphs]
    assert texts == ["1. What color is the sky?", "A) Red", "B) Blue", "C) Green"]
    assert paragraphs[2].has_any_bold() is True
    assert paragraphs[0].has_any_bold() is False


def _make_docx_with_manual_linebreaks(path):
    """Word's Shift+Enter (manual line break) keeps multiple visual lines
    inside a single <w:p> paragraph -- python-docx exposes it as an
    embedded "\\n" in Run.text rather than a paragraph boundary."""
    doc = Document()
    doc.add_paragraph("1. What color is the sky?")
    p = doc.add_paragraph()
    p.add_run("A) Red\nB) Blue\nC) Green")
    doc.save(str(path))


@pytest.mark.asyncio
async def test_manual_linebreak_splits_into_separate_lines(tmp_path):
    path = tmp_path / "quiz.docx"
    _make_docx_with_manual_linebreaks(path)

    paragraphs = await DocxParser().extract(path)

    texts = [p.plain_text for p in paragraphs]
    assert texts == ["1. What color is the sky?", "A) Red", "B) Blue", "C) Green"]


@pytest.mark.asyncio
async def test_extract_corrupted_file_raises(tmp_path):
    path = tmp_path / "not_a_docx.docx"
    path.write_bytes(b"this is not a valid docx file")

    with pytest.raises(CorruptedFileError):
        await DocxParser().extract(path)
